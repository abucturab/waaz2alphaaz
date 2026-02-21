"""Write transcribed (and optionally Alphaaz-converted) text to document formats."""

from pathlib import Path
from typing import Literal

Format = Literal["docx", "txt", "md"]


def write_document(
    content: str,
    output_path: Path,
    format: Format = "docx",
    title: str = "Transcription",
) -> None:
    """Write content to DOCX, TXT, or MD file."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if format == "txt":
        output_path = output_path if output_path.suffix == ".txt" else output_path.with_suffix(".txt")
        output_path.write_text(content, encoding="utf-8")
        return

    if format == "md":
        output_path = output_path if output_path.suffix == ".md" else output_path.with_suffix(".md")
        output_path.write_text(f"# {title}\n\n{content}", encoding="utf-8")
        return

    # docx
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise RuntimeError("python-docx is required for DOCX output: pip install python-docx")

    output_path = output_path if output_path.suffix == ".docx" else output_path.with_suffix(".docx")
    doc = Document()
    doc.add_heading(title, 0)
    for paragraph in content.split("\n\n"):
        doc.add_paragraph(paragraph)
    doc.save(str(output_path))
